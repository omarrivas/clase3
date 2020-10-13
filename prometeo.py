import speech_recognition as speechRecognition
from gtts import gTTS
from playsound import playsound
import os
from docx import Document

###################################
##MODELO DE CONTRATO
titulo= "PLIEGO DE CONDICIONES GENERALES LICITACIONES PÚBLICAS"
objeto= "Objeto: Este llamado a licitación, tiene por objeto contratar el suministro de los materiales y elementos mencionados en el detalle y especificaciones del pliego de condiciones particulares, integrante de este pliego, debiendo la mercadería satisfacer, en su calidad y características, a juicio del Poder Ejecutivo a los fines a que se destinen.-"
clausula1= "1. Tipo de Contratación."
clausula2= "2. Bien o Servicio."
clausula3= "3. Costos."
clausula4= "4. Ofertar la totalidad."
clausula5= "5. Moneda."

##################################
##MENU PRINCIPAL
saludoInicial= "Bienenido a PROMETEO. Selecciona una opción:"
opcionesMenuPrincipal= "a Bienes o Servicios b Obra Publica c Obra Publica Mayor"
pregunta1= "¿Que tipo de contratacion desea realizar? a) licitacion publica b) contratacion directa"
pregunta2= "¿Que bien o servicio deseas adquirir? (ejemplo computadoras, sillas, impresoras, limpieza)"
tiposervicio= ""
pregunta3= "¿El pliego tiene costo?"
costopliego= ""
pregunta4= "¿Se debe ofertar la totalidad del renglon?"
pregunta5= "¿Se debe ofertar por todos los renglones?"
ofertartotalidad= ""
pregunta6= "¿La oferta es en pesos argentinos?"
pregunta7= "¿La adjudicacion puede ser en moneda extranjera?"
moneda= ""

###################################
##OBTIENE RESPUESTA HABLADA
def escuchaRespuesta():
    texto= ""
    reconocedor = speechRecognition.Recognizer()
    with speechRecognition.Microphone() as fuente:
        print("Diga la respuesta: ")
        audio = reconocedor.listen(fuente)
    try:
        # texto = reconocedor.recognize_google(audio)
        texto = reconocedor.recognize_google(audio, language="es-AR")
        print("dijiste " + texto)
    except:
        print("error durante la conversion")

    texto = texto.lower()
    texto = texto.replace('á', 'a')
    texto = texto.replace('é', 'e')
    texto = texto.replace('í', 'i')
    texto = texto.replace('ó', 'o')
    texto = texto.replace('ú', 'u')
    return texto

###################################
##PRONUNCIA TEXTO
def dice(texto):
    print("PROMETEO dice: " + texto)
    tts = gTTS(texto, lang='es-us')
    tts.save("mensaje.mp3")
    playsound('mensaje.mp3')
    os.remove("mensaje.mp3")

###################################
##GENERA WORD
def generaDocumento(tiposervicio, costopliego, ofertartotalidad, moneda):
    document= Document()
    document.add_heading(titulo, level=0)
    document.add_heading(objeto, level=1)
    document.add_heading(clausula1, level=2)
    document.add_heading("Licitación Pública", level=3)
    document.add_heading(clausula2, level=2)
    document.add_heading(tiposervicio, level=3)
    document.add_heading(clausula3, level=2)
    document.add_heading(costopliego, level=3)
    document.add_heading(clausula4, level=2)
    document.add_heading(ofertartotalidad, level=3)
    document.add_heading(clausula5, level=2)
    document.add_heading(moneda, level=3)
    document.save('Pliego.docx')

#################################
##PREGUNTAS Y SETEO DEL DOCUMENTO
dice(saludoInicial)
dice(opcionesMenuPrincipal)
respuesta= escuchaRespuesta()
if respuesta== "c" or respuesta== "opcion c":
    dice(pregunta1)
    respuesta = escuchaRespuesta()
    if respuesta== "a" or respuesta== "opcion a":
        dice(pregunta2)
        respuesta = escuchaRespuesta()
        tiposervicio= respuesta

        dice(pregunta3)
        respuesta = escuchaRespuesta()
        if respuesta== "no" or respuesta== "no tiene":
            costopliego= "el pliego NO tiene costo"
        else:
            costopliego = "el pliego SI tiene costo"

        dice(pregunta4)
        respuesta = escuchaRespuesta()
        if respuesta == "no" or respuesta == "no se debe" or respuesta== "no debe":
            ofertartotalidad = "NO se debe ofertar la totalidad del renglon.\n"
        else:
            ofertartotalidad = "SI se debe ofertar la totalidad del renglon.\n"

        dice(pregunta5)
        respuesta = escuchaRespuesta()
        if respuesta == "no" or respuesta == "no se debe" or respuesta == "no debe":
            ofertartotalidad = ofertartotalidad + "NO se debe ofertar por todos los renglones."
        else:
            ofertartotalidad = ofertartotalidad + "SI se debe ofertar por todos los renglones."

        dice(pregunta6)
        respuesta = escuchaRespuesta()
        if respuesta == "no" or respuesta == "no es pesos":
            moneda = "La oferta NO es en pesos argentinos.\n"
        else:
            moneda = "La oferta SI es en pesos argentinos.\n"

        dice(pregunta7)
        respuesta = escuchaRespuesta()
        if respuesta == "no" or respuesta == "no se puede" or respuesta == "no puede":
            moneda = moneda + "NO se puede ofertar en moneda extranjera."
        else:
            moneda = moneda + "SI se puede ofertar en moneda extranjera."

        generaDocumento(tiposervicio, costopliego, ofertartotalidad, moneda)

        dice("Listo su documento fue generado. Gracias por trabajar con PROMETEO.")
    else:
        dice("Lo lamento solo puedo trabajar por ahora con Licitacion Publica. Hasta luego.")
else:
    dice("Lo lamento solo puedo trabajar por ahora con la opcion C. Hasta luego.")


----------------------------------------------------------------------------------------------------
